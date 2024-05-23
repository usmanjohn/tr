import pyautogui
import pyperclip
from docx import Document
from io import BytesIO
import time

# Function to read the document
def read_docx(file_path):
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

# Function to create a new document
def create_word_file(filename, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)

# Function to send text to ChatGPT
def send_text_to_chatgpt(text):
    pyperclip.copy(text)
    pyautogui.hotkey('ctrl', 'v')
    pyautogui.press('enter')

# Function to get the translated text
def get_translated_text():
    pyautogui.moveTo(400, 800)  # Adjust this as necessary
    pyautogui.click()
    pyautogui.hotkey('ctrl', 'a')
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(1)
    translated_text = pyperclip.paste()
    return translated_text

# Main function to handle the translation
def translate_document(doc_path, output_path):
    content = read_docx(doc_path)
    paragraphs = content.split('\n')

    # Locate ChatGPT icon on screen (adjust if necessary)
    chatgpt_icon = pyautogui.locateOnScreen('chatgpt_icon.png', confidence=0.8)
    if chatgpt_icon:
        pyautogui.moveTo(chatgpt_icon)
        pyautogui.click()
    else:
        print("ChatGPT icon not found!")
        return

    time.sleep(3)  # Wait for the chat interface to be ready

    full_translated_text = ""

    for paragraph in paragraphs:
        if paragraph.strip():
            text_to_send = f"Translate the following text to Korean:\n\n{paragraph}"
            send_text_to_chatgpt(text_to_send)
            time.sleep(10)  # Wait for the response
            translated_text = get_translated_text()
            full_translated_text += translated_text + "\n\n"

    create_word_file(output_path, full_translated_text)

# Streamlit GUI
import streamlit as st

st.title("Document Translator using ChatGPT")

uploaded_file = st.file_uploader("Upload a Word document", type="docx")

if uploaded_file is not None:
    doc_path = f"./{uploaded_file.name}"
    with open(doc_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if st.button("Translate"):
        output_path = "translated_document.docx"
        translate_document(doc_path, output_path)

        with open(output_path, "rb") as f:
            st.download_button(
                label="Download Translated Document",
                data=f,
                file_name="translated_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
