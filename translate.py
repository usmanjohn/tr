import openai
import streamlit as st
from docx import Document
from io import BytesIO

openai.api_key = st.secrets['my_key']

# Read the glossary file once
glossary_file_path = "glossaries.txt"

def read_glossary(file_path):
    glossary = {}
    try:
        with open(file_path, 'r', encoding='utf-8',errors='ignore') as f:
            for line in f:
                if line.strip():
                    term, translation = line.split('-', 1)  # Use 1 to ensure splitting only at the first hyphen
                    glossary[term.strip()] = translation.strip()
    except UnicodeDecodeError:
        st.error("Glossary file contains invalid characters. Please ensure it is UTF-8 encoded.")
    except Exception as e:
        st.error(f"An error occurred while reading the glossary file: {e}")
    return glossary

glossary = read_glossary(glossary_file_path)

def translate_document(doc, target_language="ko"):
    # System message to set the behavior of the assistant
    system_message = {
        "role": "system",
        "content": (
            f"You are a professional insurance lawyer who is good at translating documents accurately "
            f"while keeping paragraphing and styles. Use the following glossary terms for translation:\n\n" +
            "\n".join([f"{term}: {translation}" for term, translation in glossary.items()])
        )
    }
    
    # User message to provide the instruction and content
    user_message = {
        "role": "user",
        "content": f"while translating, Keep referenced document names in english for easier search. Translate the following text to {target_language}: \n\n" + "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    }
    
    messages = [system_message, user_message]

    # Send the messages to the OpenAI API
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=4000
    )

    translated_text = response.choices[0].message.content.strip()
    translated_paragraphs = translated_text.split('\n\n')

    translated_doc = Document()
    for paragraph in translated_paragraphs:
        translated_doc.add_paragraph(paragraph.strip())

    return translated_doc

st.title("Document Translator")

uploaded_file = st.file_uploader("Upload a Word document", type="docx")

if uploaded_file is not None:
    doc = Document(uploaded_file)
    language_options = [("Korean", "ko"), ("English", "en"), ("French", "fr"), ("Spanish", "es"), ("German", "de"), ("Chinese", "zh")]
    target_language = st.selectbox("Select target language", language_options, format_func=lambda x: x[0])[1]
    
    if st.button("Translate"):
        translated_doc = translate_document(doc, target_language)
        output = BytesIO()
        translated_doc.save(output)
        output.seek(0)
        st.download_button(
            label="Download Translated Document",
            data=output,
            file_name="translated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
